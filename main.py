import os
import docx
import sys
from PyPDF2 import PdfReader
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
import shutil
import io

# If modifying these scopes, delete the file token.json.
SCOPES = ['https://www.googleapis.com/auth/drive']

def search(keywords):
    print(keywords)
    # Ottieni il percorso della cartella corrente
    folder_path = "CV/"

    # Inizializza una lista vuota per i nomi dei file che contengono tutte le parole chiave
    matching_files = []

    # Scorri tutti i file nella cartella corrente
    for filename in os.listdir(folder_path):
        # Verifica se il file ha estensione .docx
        if filename.endswith(".docx"):
            # Apri il file DOCX e leggi tutto il contenuto
            doc = docx.Document(os.path.join(folder_path, filename))           
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        elif filename.endswith(".pdf"):
            # Apri il file PDF e leggi tutto il contenuto
            with open(os.path.join(folder_path, filename), "rb") as pdf_file:
                pdf_reader = PdfReader(pdf_file)
                text = "\n".join([page.extract_text() for page in pdf_reader.pages])

        # Verifica se tutte le parole chiave sono presenti nel contenuto del file
        all_keywords_present = all(keyword.lower() in text.lower() for keyword in keywords)

        if all_keywords_present:
            # Se tutte le parole chiave sono presenti, aggiungi il nome del file alla lista
            matching_files.append(filename)
    return matching_files

def searchGDrive(keywords):
    print("listdir before")
    print(os.listdir())
    print("listdir after")

    """Shows basic usage of the Drive v3 API.
    Prints the names and ids of the first 10 files the user has access to.
    """
    creds = None
    # The file token.json stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first
    # time.
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'client_secret_560342406312-8i0tv3kojidffep9sgioshchsdd7rs17_.apps.googleusercontent.com.json', SCOPES)
            creds = flow.run_local_server(port=0)
        # Save the credentials for the next run
        with open('token.json', 'w') as token:
            token.write(creds.to_json())

    try:
        service = build('drive', 'v3', credentials=creds)

        # ID della cartella da cui vuoi scaricare i file
        folder_id = '1YTdWcHYNZoX21ta4o7TsEkmuCf8zc6-B'

        # Inizializza una lista vuota per i nomi dei file che contengono tutte le parole chiave
        matching_files = []

        # Ottieni la lista dei file nella cartella
        results = service.files().list(supportsAllDrives=True, includeItemsFromAllDrives=True, q=f"'{folder_id}' in parents",
                                    fields='files(id, name, mimeType)').execute()
        files = results.get('files', [])

        # Scorri tutti i file nella cartella
        for file in files:
            file_name = file['name']
            file_id = file['id']
            file_content = download_file_content(service, file_id, file['mimeType'])

            # Verifica se tutte le parole chiave sono presenti nel contenuto del file
            all_keywords_present = all(keyword.lower() in file_content.lower() for keyword in keywords)

            if all_keywords_present:
                # Se tutte le parole chiave sono presenti, aggiungi il nome del file alla lista
                matching_files.append(file_name)
        return matching_files
        
    except HttpError as error:
        # TODO(developer) - Handle errors from drive API.
        print(f'An error occurred: {error}')
    
    
def download_file_content(service, file_id, mime_type):
    request = service.files().get_media(fileId=file_id)
    file_content = io.BytesIO()
    downloader = MediaIoBaseDownload(file_content, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    file_content.seek(0)

    if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc = docx.Document(file_content)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif mime_type == 'application/pdf':
        pdf_reader = PdfReader(file_content)
        return "\n".join([page.extract_text() for page in pdf_reader.pages])
    else:
        return ""


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Inserisci le parole chiave separate da spazi \n Usage: python3 main.py parola_chiave1 parola_chiave2 ...")
    else:
        # Ottieni le parole chiave dalla riga di comando
        keywords = sys.argv[1:]
        matching_files = search(keywords)
        #matching_files = searchGDrive(keywords)

        # Stampa i nomi dei file che contengono tutte le parole chiave
        print("I seguenti file contengono tutte le parole chiave:")
        for filename in matching_files:
            print(filename)

